#!/usr/bin/env python3
"""Build a Word document (.docx) report from report payload + chart images.

Input JSON (via stdin):
{
  "payload": { ... full report payload ... },
  "charts": { "timeSeries": "/path/to.png", ... },
  "narratives": { "executive_summary": "...", "platform_analysis": "...", ... },
  "options": { "template": "standard", "scope": "series" },
  "branding": { "company_name": "Clutch Group", ... },
  "outputPath": "/tmp/reports/{id}/report.docx"
}
"""

import sys
import os
import json
from typing import Any, Dict, List, Optional

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import (
    read_stdin_json, load_branding, hex_to_rgb,
    platform_label, tier_label,
    compact_number, format_number, format_hours, format_date, format_datetime,
    aggregate_metrics, build_scope_label,
)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Style Helpers ────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    """Set background color of a table cell."""
    r, g, b = hex_to_rgb(hex_color)
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, values: List[str], header: bool = False,
                  branding: Optional[Dict] = None):
    """Add a row to a table with optional header styling."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(str(val))
        if header and branding:
            set_cell_shading(cell, branding.get('table_header_bg', '#1e293b'))
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)
        else:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(55, 65, 81)
    return row


def add_chart_image(doc, chart_path: str, width: float = 6.5):
    """Add a chart image to the document if the file exists."""
    if chart_path and os.path.isfile(chart_path):
        doc.add_picture(chart_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def add_narrative(doc, text: str, branding: Dict):
    """Add a narrative paragraph with consistent styling."""
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*branding.get('text_color_rgb', (55, 65, 81)))
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25


def add_section_heading(doc, title: str, level: int = 2):
    """Add a styled section heading."""
    heading = doc.add_heading(title, level=level)
    return heading


# ── Document Sections ────────────────────────────────────────────────────────

def build_cover_page(doc: Document, payload: Dict, branding: Dict, scope_label: str):
    """Build the cover page with series info and branding."""
    series = payload.get('series', {})

    # Add some spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(branding.get('company_name', 'Clutch Group'))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*branding.get('muted_color_rgb', (156, 163, 175)))
    run.font.bold = True

    # Logo
    logo_path = branding.get('logo_path')
    if logo_path and os.path.isfile(logo_path):
        doc.add_picture(logo_path, width=Inches(2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Viewership Report')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*branding.get('accent_color_rgb', (59, 130, 246)))

    # Series name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(series.get('name', 'Unknown Series'))
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*branding.get('text_color_rgb', (55, 65, 81)))

    # Scope
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(scope_label)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*branding.get('muted_color_rgb', (156, 163, 175)))

    # Game + Partner
    meta_parts = []
    if series.get('game'):
        meta_parts.append(series['game'])
    if series.get('partner'):
        meta_parts.append(f"Partner: {series['partner']}")
    if meta_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(' | '.join(meta_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*branding.get('muted_color_rgb', (156, 163, 175)))

    # Dates
    start = format_date(series.get('startDate'))
    end = format_date(series.get('endDate'))
    if start != '—' or end != '—':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{start} — {end}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*branding.get('muted_color_rgb', (156, 163, 175)))

    # Generated timestamp
    gen = payload.get('generatedAt', '')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {format_datetime(gen)}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)

    doc.add_page_break()


def build_executive_summary(doc: Document, narratives: Dict, branding: Dict):
    """Build the executive summary section."""
    text = narratives.get('executive_summary', '')
    if not text:
        return
    add_section_heading(doc, 'Executive Summary', level=1)
    add_narrative(doc, text, branding)
    doc.add_paragraph()  # spacing


def build_key_metrics_table(doc: Document, agg: Dict, payload: Dict, branding: Dict):
    """Build the key metrics summary table."""
    add_section_heading(doc, 'Key Metrics', level=1)

    channels = payload.get('channels', [])
    days = payload.get('broadcastDays', [])

    metrics = [
        ('Peak Concurrent Viewers', compact_number(agg['peakCCV'])),
        ('Peak Timestamp', format_datetime(agg.get('peakTimestamp'))),
        ('Average Concurrent Viewers', compact_number(agg['avgCCV'])),
        ('Total Viewed Hours', format_number(agg['totalViewedHours'])),
        ('Tracked Channels', str(len(channels))),
        ('Broadcast Days', str(len(days))),
        ('Total Snapshots', format_number(payload.get('snapshotCount', 0))),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.add_row()
    for i, label in enumerate(['Metric', 'Value']):
        cell = hdr.cells[i]
        cell.text = label
        set_cell_shading(cell, branding.get('table_header_bg', '#1e293b'))
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(9)

    for metric_name, metric_val in metrics:
        row = table.add_row()
        row.cells[0].text = metric_name
        row.cells[1].text = metric_val
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(55, 65, 81)

    doc.add_paragraph()


def build_viewership_timeline(doc: Document, charts: Dict, narratives: Dict, branding: Dict):
    """Add the viewership timeline chart section."""
    if 'timeSeries' not in charts:
        return
    add_section_heading(doc, 'Viewership Timeline', level=1)
    add_narrative(doc, narratives.get('viewership_timeline', ''), branding)
    add_chart_image(doc, charts['timeSeries'])
    doc.add_paragraph()


def build_platform_analysis(doc: Document, charts: Dict, agg: Dict, narratives: Dict, branding: Dict):
    """Add the platform analysis section."""
    add_section_heading(doc, 'Platform Analysis', level=1)
    add_narrative(doc, narratives.get('platform_analysis', ''), branding)

    if 'platformDonut' in charts:
        add_chart_image(doc, charts['platformDonut'])

    # Platform summary table
    platforms = agg.get('platformBreakdown', [])
    if platforms:
        doc.add_paragraph()
        table = doc.add_table(rows=0, cols=4)
        table.style = 'Table Grid'

        hdr = table.add_row()
        for i, label in enumerate(['Platform', 'Total CCV', 'Avg CCV', 'Peak CCV']):
            cell = hdr.cells[i]
            cell.text = label
            set_cell_shading(cell, branding.get('table_header_bg', '#1e293b'))
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

        for p in platforms:
            row = table.add_row()
            row.cells[0].text = platform_label(p.get('platform', ''))
            row.cells[1].text = format_number(p.get('totalCCV', 0))
            row.cells[2].text = format_number(p.get('avgCCV', 0))
            row.cells[3].text = format_number(p.get('peakCCV', 0))
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()


def build_audience_breakdown(doc: Document, charts: Dict, agg: Dict, narratives: Dict, branding: Dict):
    """Add language and region breakdown sections."""
    add_section_heading(doc, 'Audience Breakdown', level=1)
    add_narrative(doc, narratives.get('audience_breakdown', ''), branding)

    if 'languageBars' in charts:
        doc.add_heading('Language Distribution', level=2)
        add_chart_image(doc, charts['languageBars'])

    if 'regionBars' in charts:
        doc.add_heading('Region Distribution', level=2)
        add_chart_image(doc, charts['regionBars'])

    doc.add_paragraph()


def build_channel_performance(doc: Document, charts: Dict, agg: Dict, branding: Dict):
    """Add the channel leaderboard section."""
    add_section_heading(doc, 'Channel Performance', level=1)

    if 'channelLeaderboard' in charts:
        add_chart_image(doc, charts['channelLeaderboard'])

    # Top channels table
    leaderboard = agg.get('channelLeaderboard', [])[:20]
    if leaderboard:
        doc.add_paragraph()
        table = doc.add_table(rows=0, cols=5)
        table.style = 'Table Grid'

        hdr = table.add_row()
        for i, label in enumerate(['#', 'Channel', 'Platform', 'Peak CCV', 'Avg CCV']):
            cell = hdr.cells[i]
            cell.text = label
            set_cell_shading(cell, branding.get('table_header_bg', '#1e293b'))
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

        for rank, ch in enumerate(leaderboard, 1):
            row = table.add_row()
            row.cells[0].text = str(rank)
            row.cells[1].text = ch.get('displayName', '')
            row.cells[2].text = platform_label(ch.get('platform', ''))
            row.cells[3].text = format_number(ch.get('peakCCV', 0))
            row.cells[4].text = format_number(ch.get('avgCCV', 0))
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()


def build_community_reach(doc: Document, payload: Dict, narratives: Dict, branding: Dict):
    """Add community reach / tier breakdown section."""
    text = narratives.get('community_reach', '')
    if not text:
        return
    add_section_heading(doc, 'Community Reach', level=1)
    add_narrative(doc, text, branding)

    # Tier breakdown table
    channels = payload.get('channels', [])
    tier_counts = {}
    for ch in channels:
        t = ch.get('tier', 'unknown')
        tier_counts[t] = tier_counts.get(t, 0) + 1

    if tier_counts:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        hdr = table.add_row()
        for i, label in enumerate(['Category', 'Channels']):
            cell = hdr.cells[i]
            cell.text = label
            set_cell_shading(cell, branding.get('table_header_bg', '#1e293b'))
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

        for tier, count in sorted(tier_counts.items(), key=lambda x: x[1], reverse=True):
            row = table.add_row()
            row.cells[0].text = tier_label(tier)
            row.cells[1].text = str(count)

    doc.add_paragraph()


def build_day_over_day(doc: Document, charts: Dict, narratives: Dict, branding: Dict):
    """Add day-over-day trend section (stage/series scope only)."""
    if 'dayOverDay' not in charts:
        return
    add_section_heading(doc, 'Day-over-Day Trend', level=1)
    add_narrative(doc, narratives.get('day_over_day', ''), branding)
    add_chart_image(doc, charts['dayOverDay'])
    doc.add_paragraph()


def build_stage_comparison(doc: Document, charts: Dict, narratives: Dict, branding: Dict):
    """Add stage comparison section (series scope only)."""
    if 'stageComparison' not in charts:
        return
    add_section_heading(doc, 'Stage Comparison', level=1)
    add_narrative(doc, narratives.get('stage_comparison', ''), branding)
    add_chart_image(doc, charts['stageComparison'])
    doc.add_paragraph()


def build_vod_metrics(doc: Document, narratives: Dict, branding: Dict):
    """Add VOD metrics section if narrative exists."""
    text = narratives.get('vod_metrics', '')
    if not text:
        return
    add_section_heading(doc, 'VOD & Clip Performance', level=1)
    add_narrative(doc, text, branding)
    doc.add_paragraph()


def build_historical_comparison(doc: Document, narratives: Dict, branding: Dict):
    """Add historical comparison section if narrative exists."""
    text = narratives.get('historical_comparison', '')
    if not text:
        return
    add_section_heading(doc, 'Historical Comparison', level=1)
    add_narrative(doc, text, branding)
    doc.add_paragraph()


def build_methodology(doc: Document, payload: Dict, branding: Dict):
    """Add methodology note section."""
    add_section_heading(doc, 'Methodology', level=1)

    scope = payload.get('scope', 'series')
    snap_count = payload.get('snapshotCount', 0)
    channels = payload.get('channels', [])

    text = (
        f"This report was generated from {format_number(snap_count)} viewership snapshots "
        f"collected across {len(channels)} tracked channels. Data is captured at regular "
        f"polling intervals during live broadcasts via platform APIs (Twitch, YouTube, Kick, TikTok). "
        f"Concurrent viewer counts (CCV) represent the number of viewers watching at each "
        f"polling interval. 'Total Viewed Hours' is calculated as the sum of concurrent viewers "
        f"across all polling intervals, converted to hours. Channel tier classifications "
        f"(Primary, Secondary, Community, Watch Party) are assigned based on viewership thresholds "
        f"and editorial judgment. Report scope: {scope}."
    )
    add_narrative(doc, text, branding)


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    input_data = read_stdin_json()
    payload = input_data['payload']
    charts = input_data.get('charts', {})
    narratives = input_data.get('narratives', {})
    options = input_data.get('options', {})
    branding = input_data.get('branding', load_branding())
    output_path = input_data['outputPath']

    scope = options.get('scope', payload.get('scope', 'series'))
    scope_label = build_scope_label(payload)
    agg = aggregate_metrics(payload.get('metrics', []))

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = branding.get('font_family', 'Arial')
    font.size = Pt(10)
    font.color.rgb = RGBColor(*branding.get('text_color_rgb', (55, 65, 81)))

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build sections in order
    build_cover_page(doc, payload, branding, scope_label)
    build_executive_summary(doc, narratives, branding)
    build_key_metrics_table(doc, agg, payload, branding)
    build_viewership_timeline(doc, charts, narratives, branding)
    build_platform_analysis(doc, charts, agg, narratives, branding)
    build_audience_breakdown(doc, charts, agg, narratives, branding)
    build_channel_performance(doc, charts, agg, branding)
    build_community_reach(doc, payload, narratives, branding)

    if scope in ('stage', 'series', 'multi_stage'):
        build_day_over_day(doc, charts, narratives, branding)
    if scope == 'series':
        build_stage_comparison(doc, charts, narratives, branding)

    build_vod_metrics(doc, narratives, branding)
    build_historical_comparison(doc, narratives, branding)
    build_methodology(doc, payload, branding)

    # Add page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(156, 163, 175)
        # Page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar2)
        run2 = p.add_run(f'  |  {branding.get("company_name", "Clutch Group")}')
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(156, 163, 175)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(json.dumps({'status': 'ok', 'path': output_path}))


if __name__ == '__main__':
    main()
